import fitz  # PyMuPDF
from docx import Document
import io
import re
import logging
from .nlp_parser import extract_contacts_nlp
from .vcard_parser import parse_vcard
from ..utils.nlp import categorize_contact

# Optional OCR imports - gracefully handle missing dependencies
try:
    import pytesseract
    from PIL import Image
    import os
    import shutil
    import subprocess

    def find_tesseract():
        """Find Tesseract executable in common locations"""
        # Try environment variable first
        env_path = os.getenv('TESSERACT_PATH')
        if env_path and os.path.isfile(env_path):
            return env_path

        # Common Tesseract paths to try
        common_paths = [
            '/usr/bin/tesseract',
            '/usr/local/bin/tesseract',
            '/opt/homebrew/bin/tesseract',  # macOS with Homebrew
            'tesseract',  # System PATH
        ]

        # Try using 'which' command
        try:
            result = subprocess.run(['which', 'tesseract'],
                                  capture_output=True, text=True, timeout=5)
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except:
            pass

        # Try shutil.which (Python's built-in)
        which_result = shutil.which('tesseract')
        if which_result:
            return which_result

        # Try common paths
        for path in common_paths:
            if path != 'tesseract' and os.path.isfile(path):
                return path

        return None

    # Find and configure Tesseract
    tesseract_path = find_tesseract()

    if tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        print(f"🔧 Found Tesseract at: {tesseract_path}")

        # Set TESSDATA_PREFIX to use bundled tessdata
        # First, try to use bundled tessdata from the application directory
        script_dir = os.path.dirname(os.path.abspath(__file__))
        bundled_tessdata = os.path.join(script_dir, '..', '..', 'tessdata')
        bundled_tessdata = os.path.abspath(bundled_tessdata)

        tessdata_prefix = None

        # Check if bundled tessdata exists and has required files
        if os.path.isdir(bundled_tessdata) and os.path.isfile(os.path.join(bundled_tessdata, 'eng.traineddata')):
            tessdata_prefix = bundled_tessdata
            print(f"✅ Using bundled tessdata: {tessdata_prefix}")
        else:
            # Fallback to environment variable or system paths
            tessdata_prefix = os.getenv('TESSDATA_PREFIX')
            if not tessdata_prefix:
                # Try to find tessdata directory automatically
                common_paths = [
                    '/usr/share/tesseract-ocr/tessdata',
                    '/usr/share/tesseract-ocr/4.00/tessdata',
                    '/usr/share/tesseract-ocr/5.00/tessdata',
                    '/usr/share/tessdata',
                    '/usr/local/share/tessdata'
                ]

                for path in common_paths:
                    if os.path.isdir(path) and os.path.isfile(os.path.join(path, 'eng.traineddata')):
                        tessdata_prefix = path
                        print(f"🔍 Auto-detected system tessdata: {tessdata_prefix}")
                        break

        if tessdata_prefix:
            os.environ['TESSDATA_PREFIX'] = tessdata_prefix
            print(f"🔧 Set TESSDATA_PREFIX: {tessdata_prefix}")

            # Verify the eng.traineddata file exists
            eng_file = os.path.join(tessdata_prefix, 'eng.traineddata')
            if os.path.isfile(eng_file):
                print(f"✅ Found English language data: {eng_file}")
                # List all available language files
                try:
                    lang_files = [f for f in os.listdir(tessdata_prefix) if f.endswith('.traineddata')]
                    print(f"📋 Available languages: {', '.join([f.replace('.traineddata', '') for f in lang_files])}")
                except:
                    pass
            else:
                print(f"⚠️  English language data not found at: {eng_file}")
        else:
            print("⚠️  TESSDATA_PREFIX not set and could not auto-detect")

        # Test if tesseract is actually working
        try:
            version = pytesseract.get_tesseract_version()
            OCR_AVAILABLE = True
            print(f"✅ Tesseract OCR is available: {version}")
            if tessdata_prefix:
                print(f"✅ Using Tesseract data from: {tessdata_prefix}")
        except Exception as e:
            OCR_AVAILABLE = False
            print(f"⚠️  Tesseract found but not working: {e}")
            if tessdata_prefix:
                print(f"   Check TESSDATA_PREFIX: {tessdata_prefix}")
    else:
        OCR_AVAILABLE = False
        print("⚠️  Tesseract not found in common locations")
        print("   Checked paths: /usr/bin/tesseract, /usr/local/bin/tesseract, system PATH")
        print("   OCR functionality will be disabled")

except ImportError as e:
    OCR_AVAILABLE = False
    print(f"⚠️  OCR dependencies not available: {e}")
    print("   Install with: pip install pytesseract pillow")
    print("   Image parsing will be disabled")

logger = logging.getLogger(__name__)

def parse_pdf(file_content):
    doc = fitz.open(stream=file_content, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return extract_contacts(text)

def parse_docx(file_content):
    doc = Document(io.BytesIO(file_content))
    text = "\n".join([para.text for para in doc.paragraphs])
    return extract_contacts(text)

def parse_txt(file_content):
    text = file_content.decode("utf-8")

    # Check if this is structured data (tab-separated or CSV-like)
    lines = text.strip().split('\n')
    if len(lines) > 1:
        # Check if first line looks like headers
        first_line = lines[0].lower()
        if any(header in first_line for header in ['name', 'email', 'phone', 'company', 'designation']):
            # Detect delimiter
            if '\t' in lines[0]:
                return parse_structured_text(text, delimiter='\t')
            elif ',' in lines[0] and lines[0].count(',') > 2:
                return parse_structured_text(text, delimiter=',')

    # Fallback to regular text parsing
    return extract_contacts(text)

def parse_structured_text(text, delimiter='\t'):
    """Parse structured text data (tab-separated or CSV-like)"""
    lines = text.strip().split('\n')
    if len(lines) < 2:
        return []

    # Parse header row
    headers = [h.strip().lower() for h in lines[0].split(delimiter)]
    contacts = []

    # Create field mapping
    field_mapping = {
        'name': ['name', 'full name', 'contact name'],
        'designation': ['designation', 'title', 'position', 'job title'],
        'company': ['company', 'organization', 'org', 'business'],
        'phone': ['phone', 'telephone', 'tel', 'mobile', 'cell'],
        'email': ['email', 'e-mail', 'mail'],
        'website': ['website', 'web', 'url', 'site'],
        'category': ['category', 'type', 'classification'],
        'address': ['address', 'location', 'addr'],
        'notes': ['notes', 'note', 'comments', 'remarks', 'description']
    }

    # Find column indices for each field
    column_map = {}
    for field, possible_names in field_mapping.items():
        for i, header in enumerate(headers):
            if any(name in header for name in possible_names):
                column_map[field] = i
                break

    # Parse data rows
    for line_num, line in enumerate(lines[1:], 2):
        if not line.strip():
            continue

        try:
            values = line.split(delimiter)
            contact = {
                'name': '',
                'designation': '',
                'company': '',
                'phone': '',
                'email': '',
                'website': '',
                'category': 'Others',
                'address': '',
                'notes': ''
            }

            # Extract values based on column mapping
            for field, col_index in column_map.items():
                if col_index < len(values):
                    value = values[col_index].strip()
                    if value:
                        contact[field] = value

            # Apply intelligent categorization if category is not provided or is generic
            if not contact['category'] or contact['category'].lower() in ['others', 'uncategorized', '']:
                contact['category'] = categorize_contact(contact)

            # Only add contact if it has at least a name or email
            if contact['name'] or contact['email']:
                contacts.append(contact)

        except Exception as e:
            logger.warning(f"Error parsing line {line_num}: {e}")
            continue

    return contacts

def parse_vcf(file_content):
    """Parse vCard (.vcf) files"""
    try:
        text = file_content.decode("utf-8")
        logger.info("Parsing vCard file")
        return parse_vcard(text)
    except Exception as e:
        logger.error(f"Error parsing vCard file: {e}")
        return []

def parse_image(file_content):
    """Enhanced image parsing with preprocessing and multiple OCR strategies"""
    if not OCR_AVAILABLE:
        logger.warning("OCR not available. Cannot parse image files.")
        return []

    try:
        # Load and preprocess image
        image = Image.open(io.BytesIO(file_content))

        # Try multiple OCR strategies for better accuracy
        extracted_texts = []

        # Strategy 1: Original image with default settings
        try:
            text1 = pytesseract.image_to_string(image, config='--psm 6')
            if text1.strip():
                extracted_texts.append(("default", text1))
        except:
            pass

        # Strategy 2: Preprocessed image for better OCR
        try:
            processed_image = preprocess_business_card_image(image)
            text2 = pytesseract.image_to_string(processed_image, config='--psm 6')
            if text2.strip():
                extracted_texts.append(("preprocessed", text2))
        except:
            pass

        # Strategy 3: Different PSM modes for complex layouts
        try:
            text3 = pytesseract.image_to_string(image, config='--psm 4')  # Single column text
            if text3.strip():
                extracted_texts.append(("single_column", text3))
        except:
            pass

        # Strategy 4: OCR with data extraction
        try:
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            text4 = extract_text_with_confidence(data)
            if text4.strip():
                extracted_texts.append(("high_confidence", text4))
        except:
            pass

        # Combine and process all extracted texts
        all_contacts = []
        for strategy, text in extracted_texts:
            logger.info(f"OCR Strategy '{strategy}' extracted {len(text)} characters")
            contacts = extract_contacts_advanced(text, strategy)
            all_contacts.extend(contacts)

        # Deduplicate and merge contacts
        merged_contacts = merge_duplicate_contacts(all_contacts)

        logger.info(f"Final result: {len(merged_contacts)} contacts after processing {len(extracted_texts)} OCR strategies")
        return merged_contacts

    except Exception as e:
        logger.error(f"Error parsing image: {e}")
        return []

def parse_image_fast(file_content):
    """Ultra-fast image parsing optimized for large files and timeout-sensitive operations"""
    if not OCR_AVAILABLE:
        logger.warning("OCR not available. Cannot parse image files.")
        return []

    try:
        # Load image and check size
        image = Image.open(io.BytesIO(file_content))
        original_size = image.size
        file_size_mb = len(file_content) / (1024 * 1024)
        logger.info(f"Processing image: {original_size} pixels, {file_size_mb:.1f}MB, mode: {image.mode}")

        # Render-optimized OCR configuration
        is_render = os.getenv("ENVIRONMENT") == "production"

        if is_render:
            # Render deployment: use most aggressive optimization
            logger.info("Render deployment detected, using maximum optimization")
            processed_image = preprocess_business_card_image(image, file_size_mb)
            # Use fastest OCR settings for Render
            ocr_config = '--psm 6 --oem 1 -c tessedit_do_invert=0'  # Legacy engine, no inversion
        elif file_size_mb > 1.0:  # Files larger than 1MB
            logger.info("Large file detected, using aggressive optimization")
            processed_image = preprocess_business_card_image(image, file_size_mb)
            # Use faster OCR settings for large files
            ocr_config = '--psm 6 --oem 1'  # Use legacy OCR engine for speed
        else:
            # For smaller files, use standard preprocessing
            processed_image = preprocess_business_card_image(image, file_size_mb)
            ocr_config = '--psm 6 --oem 3'  # Use default OCR engine

        try:
            # Single OCR attempt with optimized settings
            text = pytesseract.image_to_string(processed_image, config=ocr_config)

            logger.info(f"OCR extracted {len(text)} characters")

            if text.strip():
                contacts = extract_contacts_advanced(text, "fast")
                logger.info(f"Extracted {len(contacts)} contacts from OCR text")
                return contacts
            else:
                logger.warning("No text extracted from image")
                return []

        except Exception as ocr_error:
            logger.error(f"OCR processing failed: {ocr_error}")
            return []

    except Exception as e:
        logger.error(f"Error in fast image parsing: {e}")
        return []

def preprocess_business_card_image(image, file_size_mb=0):
    """Ultra-fast preprocessing optimized for Render deployment and large files"""
    try:
        # Check if running on Render for maximum optimization
        is_render = os.getenv("ENVIRONMENT") == "production"

        # For large images, resize first to reduce processing time
        width, height = image.size

        # More aggressive resizing for Render deployment
        if is_render:
            # Even more aggressive for Render - prioritize speed over quality
            if file_size_mb > 1.0:
                max_dimension = 600  # Very small for large files on Render
            else:
                max_dimension = 800  # Small for Render to ensure speed
        else:
            max_dimension = 1200  # Standard size for local/other deployments

        # Resize large images first to speed up processing
        if width > max_dimension or height > max_dimension:
            scale_factor = min(max_dimension/width, max_dimension/height)
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            # Use NEAREST for maximum speed
            image = image.resize((new_width, new_height), Image.Resampling.NEAREST)
            logger.info(f"Resized image from {width}x{height} to {new_width}x{new_height} (Render: {is_render})")

        # Convert to grayscale immediately for speed
        if image.mode != 'L':
            gray_image = image.convert('L')
        else:
            gray_image = image

        # Skip enhancement for Render to save processing time
        if is_render:
            logger.info("Render deployment: skipping image enhancement for speed")
            return gray_image
        else:
            # Minimal processing for speed - just basic contrast
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(gray_image)
            enhanced_image = enhancer.enhance(1.1)  # Very light enhancement
            return enhanced_image

    except Exception as e:
        logger.warning(f"Image preprocessing failed: {e}, using original image")
        # Return grayscale version of original if preprocessing fails
        try:
            return image.convert('L') if image.mode != 'L' else image
        except:
            return image

def extract_text_with_confidence(ocr_data):
    """Extract text from OCR data with confidence filtering"""
    try:
        texts = []
        confidences = ocr_data['conf']
        words = ocr_data['text']

        for i, (word, conf) in enumerate(zip(words, confidences)):
            # Only include words with confidence > 30
            if conf > 30 and word.strip():
                texts.append(word)

        return ' '.join(texts)
    except Exception as e:
        logger.warning(f"Confidence filtering failed: {e}")
        return ''

def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        return ""

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content"""
    try:
        from docx import Document
        import io
        doc = Document(io.BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        return ""

def parse_csv_content(content: bytes) -> List[Dict[str, str]]:
    """Parse CSV content and return list of contacts"""
    try:
        import csv
        import io

        content_str = content.decode('utf-8', errors='ignore')
        csv_reader = csv.DictReader(io.StringIO(content_str))

        contacts = []
        for row in csv_reader:
            # Map common CSV column names to our schema
            contact = {
                "name": row.get("name", row.get("Name", row.get("NAME", ""))),
                "designation": row.get("designation", row.get("Designation", row.get("title", row.get("Title", "")))),
                "company": row.get("company", row.get("Company", row.get("organization", row.get("Organization", "")))),
                "email": row.get("email", row.get("Email", row.get("EMAIL", ""))),
                "phone": row.get("phone", row.get("Phone", row.get("telephone", row.get("Telephone", "")))),
                "website": row.get("website", row.get("Website", row.get("url", row.get("URL", "")))),
                "address": row.get("address", row.get("Address", row.get("location", row.get("Location", "")))),
                "categories": "Others"
            }
            contacts.append(contact)

        return contacts
    except Exception as e:
        logger.error(f"CSV parsing failed: {e}")
        return []

def extract_contacts_advanced(text, strategy="default"):
    """Advanced contact extraction with better pattern recognition for business cards"""
    contacts = []

    # Clean and normalize text
    text = clean_ocr_text(text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]

    if not lines:
        return contacts

    # Enhanced regex patterns for business cards
    patterns = {
        'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        'phone': r'(?:(?:\+|00)[1-9]\d{0,3}[-.\s]?)?(?:\(?0?\d{1,4}\)?[-.\s]?)?\d{3,4}[-.\s]?\d{3,4}(?:[-.\s]?\d{1,4})?',
        'mobile': r'(?:\+91|91)?[-.\s]?[6-9]\d{9}',
        'website': r'(?:https?://)?(?:www\.)?[a-zA-Z0-9](?:[a-zA-Z0-9-]{0,61}[a-zA-Z0-9])?(?:\.[a-zA-Z]{2,})+',
        'designation': r'\b(?:CEO|CTO|CFO|Manager|Director|President|Vice President|VP|Senior|Junior|Executive|Officer|Head|Lead|Coordinator|Specialist|Analyst|Engineer|Developer|Designer|Consultant|Advisor)\b',
        'company_indicators': r'\b(?:Ltd|Limited|Inc|Corporation|Corp|LLC|LLP|Pvt|Private|Company|Co|Group|Associates|Partners|Solutions|Services|Technologies|Tech|Systems|Enterprises|Industries)\b'
    }

    # Try to identify the structure of the business card
    contact_info = {
        'name': '',
        'designation': '',
        'company': '',
        'email': '',
        'phone': '',
        'website': '',
        'address': '',
        'category': 'Business',
        'notes': ''
    }

    # Process lines to extract information
    address_lines = []
    notes_lines = []

    for i, line in enumerate(lines):
        line_lower = line.lower()

        # Extract email
        email_match = re.search(patterns['email'], line, re.IGNORECASE)
        if email_match and not contact_info['email']:
            contact_info['email'] = email_match.group().lower()
            continue

        # Extract phone numbers
        phone_match = re.search(patterns['phone'], line)
        if phone_match and not contact_info['phone']:
            phone = re.sub(r'[^\d+]', '', phone_match.group())
            if len(phone) >= 10:
                contact_info['phone'] = phone
                continue

        # Extract website
        website_match = re.search(patterns['website'], line, re.IGNORECASE)
        if website_match and not contact_info['website']:
            website = website_match.group()
            if not website.startswith('http'):
                website = 'https://' + website
            contact_info['website'] = website
            continue

        # Check for designation keywords
        if re.search(patterns['designation'], line, re.IGNORECASE) and not contact_info['designation']:
            contact_info['designation'] = line.strip()
            continue

        # Check for company indicators
        if re.search(patterns['company_indicators'], line, re.IGNORECASE) and not contact_info['company']:
            contact_info['company'] = line.strip()
            continue

        # Identify name (usually first non-empty line or line without special patterns)
        if (not contact_info['name'] and i < 3 and
            not re.search(patterns['email'], line) and
            not re.search(patterns['phone'], line) and
            not re.search(patterns['website'], line) and
            len(line.split()) <= 4 and
            not any(char.isdigit() for char in line)):
            contact_info['name'] = line.strip()
            continue

        # Check if line looks like an address
        if is_address_line(line):
            address_lines.append(line)
        else:
            notes_lines.append(line)

    # Post-process extracted information
    contact_info['address'] = ' '.join(address_lines)
    contact_info['notes'] = ' '.join(notes_lines)

    # Smart fallbacks and improvements
    contact_info = improve_contact_info(contact_info, lines)

    # Only add contact if we have at least name or email
    if contact_info['name'] or contact_info['email']:
        contacts.append(contact_info)

    return contacts

def clean_ocr_text(text):
    """Clean and normalize OCR text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)

    # Fix common OCR errors
    replacements = {
        r'@': '@',  # Ensure @ symbol is correct
        r'[|]': 'I',  # Common OCR mistake
        r'0(?=[A-Za-z])': 'O',  # Zero to O in words
        r'(?<=[A-Za-z])0': 'O',  # Zero to O in words
        r'5(?=[A-Za-z])': 'S',  # 5 to S in words
        r'1(?=[A-Za-z])': 'I',  # 1 to I in words
    }

    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text)

    return text.strip()

def is_address_line(line):
    """Check if a line looks like part of an address"""
    address_keywords = [
        'street', 'st', 'road', 'rd', 'avenue', 'ave', 'lane', 'ln', 'drive', 'dr',
        'building', 'bldg', 'floor', 'fl', 'suite', 'ste', 'apartment', 'apt',
        'house', 'complex', 'center', 'centre', 'plaza', 'tower', 'block',
        'near', 'opp', 'opposite', 'behind', 'beside', 'next to',
        'city', 'state', 'country', 'pin', 'zip', 'postal'
    ]

    line_lower = line.lower()

    # Check for address keywords
    if any(keyword in line_lower for keyword in address_keywords):
        return True

    # Check for patterns that look like addresses
    if re.search(r'\d+.*(?:street|road|avenue|lane|drive)', line_lower):
        return True

    # Check for postal codes
    if re.search(r'\b\d{5,6}\b', line):
        return True

    return False

def improve_contact_info(contact_info, lines):
    """Improve contact information using context and smart guessing"""

    # If no name found, try to guess from first meaningful line
    if not contact_info['name']:
        for line in lines[:3]:  # Check first 3 lines
            if (not re.search(r'[@.]', line) and  # No email or website
                not re.search(r'\d{3,}', line) and  # No long numbers
                len(line.split()) <= 4 and  # Not too many words
                line.strip()):
                contact_info['name'] = line.strip()
                break

    # If no company found, look for lines with company indicators
    if not contact_info['company']:
        company_patterns = [
            r'\b(?:Ltd|Limited|Inc|Corporation|Corp|LLC|LLP|Pvt|Private|Company|Co)\b',
            r'\b(?:Group|Associates|Partners|Solutions|Services|Technologies|Tech)\b',
            r'\b(?:Systems|Enterprises|Industries|International|Global)\b'
        ]

        for line in lines:
            for pattern in company_patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    contact_info['company'] = line.strip()
                    break
            if contact_info['company']:
                break

    # Clean up phone number
    if contact_info['phone']:
        phone = re.sub(r'[^\d+]', '', contact_info['phone'])
        if phone.startswith('+91'):
            phone = phone[3:]
        elif phone.startswith('91') and len(phone) > 10:
            phone = phone[2:]
        contact_info['phone'] = phone[-10:] if len(phone) > 10 else phone

    # Categorize based on content
    if contact_info['company'] or contact_info['designation']:
        contact_info['category'] = 'Business'
    elif contact_info['email'] and any(domain in contact_info['email'] for domain in ['.gov', '.edu']):
        contact_info['category'] = 'Government' if '.gov' in contact_info['email'] else 'Education'
    else:
        contact_info['category'] = 'Others'

    return contact_info

def merge_duplicate_contacts(contacts):
    """Merge duplicate contacts from different OCR strategies"""
    if not contacts:
        return contacts

    merged = []

    for contact in contacts:
        # Find if this contact already exists
        existing = None
        for merged_contact in merged:
            if (contact['email'] and contact['email'] == merged_contact['email']) or \
               (contact['phone'] and contact['phone'] == merged_contact['phone']) or \
               (contact['name'] and contact['name'] == merged_contact['name']):
                existing = merged_contact
                break

        if existing:
            # Merge information, preferring non-empty values
            for key in contact:
                if contact[key] and not existing[key]:
                    existing[key] = contact[key]
                elif contact[key] and len(contact[key]) > len(existing[key]):
                    existing[key] = contact[key]
        else:
            merged.append(contact)

    return merged

def extract_contacts_icc_format(text):
    """Specialized parser for ICC contact data format"""
    contacts = []
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # Enhanced patterns for ICC data
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(?:Tel|Phone|Fax|Mobile|Cell)?\s*(?:\+?91\s*)?(?:0\d{2,4}\s*)?(\d{8,12})'

    # Group lines by email (each email represents a new contact)
    contact_groups = []
    current_group = []

    for line in lines:
        if re.search(email_pattern, line):
            if current_group:
                contact_groups.append(current_group)
            current_group = [line]
        else:
            current_group.append(line)

    if current_group:
        contact_groups.append(current_group)

    # Process each group
    for group in contact_groups:
        contact = {
            "name": "",
            "email": "",
            "phone": "",
            "address": "",
            "category": "Work",  # ICC contacts are business contacts
            "notes": ""
        }

        tel_numbers = []
        fax_numbers = []
        address_parts = []

        for line in group:
            # Extract email
            emails = re.findall(email_pattern, line)
            if emails:
                contact["email"] = emails[0]
                # Extract city/location from email
                if 'icc_' in emails[0]:
                    city_code = emails[0].split('icc_')[1].split('@')[0]
                    city_mapping = {
                        'amd': 'Ahmedabad',
                        'chennai': 'Chennai',
                        'pune': 'Pune',
                        'baroda': 'Baroda/Vadodara',
                        'cjb': 'Coimbatore',
                        'delhi': 'Delhi',
                        'erode': 'Erode',
                        'hyd': 'Hyderabad',
                        'jaipur': 'Jaipur',
                        'karur': 'Karur',
                        'kolkata': 'Kolkata',
                        'tirupur': 'Tirupur',
                        'kanpur': 'Kanpur'
                    }
                    city = city_mapping.get(city_code, city_code.title())
                    contact["name"] = f"ICC {city}"
                continue

            # Extract phone numbers
            if line.lower().startswith('tel'):
                phones = re.findall(r'[\d\s]+', line)
                for phone in phones:
                    clean_phone = re.sub(r'\D', '', phone)
                    if len(clean_phone) >= 8:
                        tel_numbers.append(clean_phone)
                continue

            # Extract fax numbers
            if line.lower().startswith('fax'):
                phones = re.findall(r'[\d\s]+', line)
                for phone in phones:
                    clean_phone = re.sub(r'\D', '', phone)
                    if len(clean_phone) >= 8:
                        fax_numbers.append(clean_phone)
                continue

            # Address components
            if (any(keyword in line.lower() for keyword in ['road', 'street', 'lane', 'building', 'floor', 'complex', 'center', 'house', 'near']) or
                re.match(r'^[A-Z]?/?[\d-]+[A-Z]?$', line.strip()) or
                any(city in line for city in ['India', 'Mumbai', 'Delhi', 'Bangalore', 'Chennai', 'Kolkata', 'Pune', 'Hyderabad', 'Ahmedabad'])):
                address_parts.append(line)

        # Set phone (prefer tel over fax)
        if tel_numbers:
            contact["phone"] = tel_numbers[0]
        elif fax_numbers:
            contact["phone"] = fax_numbers[0]

        # Set address
        contact["address"] = ", ".join(address_parts) if address_parts else ""

        # Add fax to notes if different from phone
        if fax_numbers and (not tel_numbers or fax_numbers[0] != tel_numbers[0]):
            contact["notes"] = f"Fax: {fax_numbers[0]}"

        # Only add if we have meaningful data
        if contact["email"] or contact["phone"]:
            if not contact["name"]:
                contact["name"] = "ICC Office"
            contacts.append(contact)

    return contacts

def extract_contacts(text):
    """Enhanced contact extraction with intelligent grouping and pattern matching"""
    # Check if this looks like ICC format data
    if 'icc_' in text.lower() and '@iccworld.com' in text.lower():
        return extract_contacts_icc_format(text)

    # Try advanced extraction first (better for business cards)
    try:
        advanced_contacts = extract_contacts_advanced(text)
        if advanced_contacts and len(advanced_contacts) > 0:
            logger.info(f"Advanced parser extracted {len(advanced_contacts)} contacts")
            return advanced_contacts
    except Exception as e:
        logger.warning(f"Advanced parsing failed: {e}, trying NLP parser")

    # Try NLP-based extraction for unstructured text
    try:
        nlp_contacts = extract_contacts_nlp(text)
        if nlp_contacts and len(nlp_contacts) > 0:
            logger.info(f"NLP parser extracted {len(nlp_contacts)} contacts")
            return nlp_contacts
    except Exception as e:
        logger.warning(f"NLP parsing failed: {e}, falling back to regex parser")

    # Original general parser for other formats
    contacts = []
    lines = [line.strip() for line in text.split("\n") if line.strip()]

    # Enhanced patterns
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'(?:Tel|Phone|Fax|Mobile|Cell)?\s*(?:\+?91\s*)?(?:0\d{2,4}\s*)?(\d{8,12})'
    address_keywords = ['road', 'street', 'avenue', 'lane', 'building', 'floor', 'suite', 'apartment', 'house', 'complex', 'center', 'plaza', 'tower', 'near', 'opp', 'opposite']

    # Group related lines together
    contact_groups = []
    current_group = []

    for i, line in enumerate(lines):
        # Check if this line starts a new contact group
        if (re.search(email_pattern, line) or
            any(keyword in line.lower() for keyword in ['email', 'tel', 'phone', 'fax']) or
            (i > 0 and line and not line[0].isdigit() and not line.startswith(('Tel', 'Fax', 'Email')))):

            if current_group:
                contact_groups.append(current_group)
                current_group = []

        current_group.append(line)

    if current_group:
        contact_groups.append(current_group)

    # Process each group to extract contact information
    for group in contact_groups:
        contact = {
            "name": "",
            "email": "",
            "phone": "",
            "address": "",
            "category": "Uncategorized",
            "notes": ""
        }

        address_parts = []
        notes_parts = []

        for line in group:
            # Extract email
            emails = re.findall(email_pattern, line)
            if emails:
                contact["email"] = emails[0]
                continue

            # Extract phone numbers
            phones = re.findall(phone_pattern, line)
            if phones and not contact["phone"]:
                # Clean and format phone number
                phone = re.sub(r'\D', '', phones[0])
                if len(phone) >= 8:
                    contact["phone"] = phone[-10:] if len(phone) > 10 else phone
                continue

            # Check if line contains address information
            if any(keyword in line.lower() for keyword in address_keywords):
                address_parts.append(line)
                continue

            # Check for location/city information
            if any(keyword in line.lower() for keyword in ['india', 'mumbai', 'delhi', 'bangalore', 'chennai', 'kolkata', 'pune', 'hyderabad']):
                address_parts.append(line)
                continue

            # If line contains building/office identifiers
            if re.match(r'^[A-Z]?/?[\d-]+[A-Z]?$', line.strip()) or line.strip().endswith(('Floor', 'floor')):
                address_parts.append(line)
                continue

            # Extract name (usually the first meaningful line that's not tel/email/address)
            if (not contact["name"] and
                not line.lower().startswith(('tel', 'fax', 'email', 'phone')) and
                not re.match(r'^\d+$', line.strip()) and
                len(line.split()) >= 2):
                contact["name"] = line.strip()
                continue

            # Everything else goes to notes
            if line and not re.match(r'^\d+$', line.strip()):
                notes_parts.append(line)

        # Compile address and notes
        contact["address"] = ", ".join(address_parts) if address_parts else ""
        contact["notes"] = ", ".join(notes_parts) if notes_parts else ""

        # Only add contact if we have at least a name, email, or phone
        if contact["name"] or contact["email"] or contact["phone"]:
            # If no name but have email, extract name from email
            if not contact["name"] and contact["email"]:
                email_name = contact["email"].split('@')[0].replace('.', ' ').replace('_', ' ')
                contact["name"] = email_name.title()

            # If still no name, create one from available info
            if not contact["name"]:
                if contact["phone"]:
                    contact["name"] = f"Contact {contact['phone'][-4:]}"
                else:
                    contact["name"] = "Unknown Contact"

            contacts.append(contact)

    return contacts